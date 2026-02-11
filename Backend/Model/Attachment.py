from pydantic import BaseModel, Field
from docx import Document
import os


class Attachment(BaseModel):
    application_code: str = Field(..., description="Application code for the letter")
    date: str = Field(..., description="Date of the letter")
    prefix: str = Field(..., description="Prefix (e.g. Mr., Ms., Dr.)")
    last_name: str = Field(..., description="Last name of the recipient")
    street: str = Field(..., description="Street address of the recipient")
    city: str = Field(..., description="City of the recipient")
    name: str = Field(..., description="Full name of the recipient (without last name)")
    position: str = Field(..., description="Position applied for")

    education_required: str = ""
    education: str = ""
    experience_required: str = ""
    experience: str = ""
    training_required: str = ""
    training: str = ""
    eligibility_required: str = ""
    eligibility: str = ""
    education_remarks: str = ""
    experience_remarks: str = ""
    training_remarks: str = ""
    eligibility_remarks: str = ""
    performance_required: str = ""
    performance: str = ""
    remarks: str = ""


    def generate_docx(self, template_path: str, output_path: str):
        try:
            doc = Document(template_path)
            replacements = {
                "[Application Code]": self.application_code,
                "[Date]": self.date,
                "[Prefix]": self.prefix,
                "[LastName]": self.last_name,
                "[Street]": self.street,
                "[City]": self.city,
                "[Name]": f"{self.name} {self.last_name}",
                "[Position]": self.position,
                "[Education Required]": self.education_required,
                "[Education]": self.education,
                "[Experience Required]": self.experience_required,
                "[Experience]": self.experience,
                "[Training Required]": self.training_required,
                "[Training]": self.training,
                "[Eligibility Required]": self.eligibility_required,
                "[Eligibility]": self.eligibility,
                "[Education Remarks]": self.education_remarks,
                "[Experience Remarks]": self.experience_remarks,
                "[Training Remarks]": self.training_remarks,
                "[Eligibility Remarks]": self.eligibility_remarks,
                "[Performance Required]": self.performance_required,
                "[Performance]": self.performance,
                "[Remarks]": self.remarks,
            }

            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, value in replacements.items():
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, value)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)

            doc.save(output_path)

            return {"message": "Doc generated", "statusCode": 200}
        except Exception as e:
            return {"statusCode": 500, "message": str(e)}
